"""Knowledge loaders with privacy-safe source identifiers."""

from __future__ import annotations

from hashlib import sha256
from io import BytesIO
from pathlib import Path
from typing import BinaryIO

from gamesight.knowledge.models import KnowledgeDocument


SUPPORTED_KNOWLEDGE_SUFFIXES = {".md", ".txt", ".docx"}


def _stable_id(title: str, content: str) -> str:
    digest = sha256(f"{title}\n{content}".encode("utf-8")).hexdigest()[:20]
    return f"doc_{digest}"


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _docx_text(source: str | Path | BinaryIO) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - optional dependency path
        raise ImportError(
            "DOCX knowledge sources require python-docx. "
            "Install with: pip install python-docx"
        ) from exc
    document = Document(source)
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = (paragraph.style.name if paragraph.style else "").lower()
        if style.startswith("heading"):
            level = "".join(character for character in style if character.isdigit())
            prefix = "#" * max(1, min(6, int(level or "2")))
            blocks.append(f"{prefix} {text}")
        else:
            blocks.append(text)
    for table in document.tables:
        rows = []
        for row in table.rows:
            values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(values):
                rows.append(" | ".join(values))
        if rows:
            blocks.append("\n".join(rows))
    return "\n\n".join(blocks)


def _title_from_content(filename: str, content: str) -> str:
    for line in content.splitlines():
        candidate = line.strip().lstrip("#").strip()
        if candidate:
            return candidate[:160]
    return Path(filename).stem


def load_knowledge_bytes(
    data: bytes,
    filename: str,
    *,
    source_uri: str | None = None,
    language: str | None = None,
    metadata: dict | None = None,
) -> KnowledgeDocument:
    """Load an uploaded source without exposing a machine-local absolute path."""
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_KNOWLEDGE_SUFFIXES:
        raise ValueError(f"Unsupported knowledge source: {suffix or filename}")
    content = _docx_text(BytesIO(data)) if suffix == ".docx" else _decode_text(data)
    content = content.strip()
    if not content:
        raise ValueError(f"Knowledge source is empty: {filename}")
    title = _title_from_content(filename, content)
    return KnowledgeDocument(
        document_id=_stable_id(title, content),
        title=title,
        source_uri=source_uri or f"upload://{Path(filename).name}",
        content=content,
        language=language,
        metadata={"source_name": Path(filename).name, **(metadata or {})},
    )


def load_knowledge_document(
    path: str | Path,
    *,
    source_uri: str | None = None,
    language: str | None = None,
    metadata: dict | None = None,
) -> KnowledgeDocument:
    """Load Markdown, text or DOCX while keeping absolute paths out of outputs."""
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_KNOWLEDGE_SUFFIXES:
        raise ValueError(f"Unsupported knowledge source: {suffix or path.name}")
    if suffix == ".docx":
        content = _docx_text(path)
    else:
        content = _decode_text(path.read_bytes())
    content = content.strip()
    if not content:
        raise ValueError(f"Knowledge source is empty: {path.name}")
    title = _title_from_content(path.name, content)
    return KnowledgeDocument(
        document_id=_stable_id(title, content),
        title=title,
        source_uri=source_uri or f"local://{path.name}",
        content=content,
        language=language,
        metadata={"source_name": path.name, **(metadata or {})},
    )
