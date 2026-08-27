"""Tests for deterministic knowledge ingestion and retrieval."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory

from gamesight.knowledge.chunking import KnowledgeChunker
from gamesight.knowledge.embeddings import EmbeddingProvider
from gamesight.knowledge.dynamic_data import dynamic_game_documents
from gamesight.knowledge.indexing import index_documents
from gamesight.knowledge.loaders import load_knowledge_bytes, load_knowledge_document
from gamesight.knowledge.models import (
    KnowledgeChunk,
    KnowledgeDocument,
    KnowledgeLayer,
    RuleStrength,
)
from gamesight.knowledge.retriever import KnowledgeRetriever
from gamesight.knowledge.routing import DecisionContext
from gamesight.knowledge.store import (
    ChromaKnowledgeStore,
    InMemoryKnowledgeStore,
    LayeredKnowledgeStore,
)


class KeywordEmbedding(EmbeddingProvider):
    vocabulary = ("flash", "cover", "economy", "地图", "掩体")

    @property
    def model_id(self) -> str:
        return "test-keywords"

    def _embed(self, text: str) -> list[float]:
        lowered = text.lower()
        return [float(lowered.count(word)) for word in self.vocabulary]

    def embed_documents(self, texts):
        return [self._embed(text) for text in texts]

    def embed_query(self, text):
        return self._embed(text)


def test_load_text_bytes_uses_private_upload_uri() -> None:
    document = load_knowledge_bytes(
        "# 闪光弹\n\n看到闪光时应利用掩体。".encode("utf-8"),
        "cs2-guide.md",
        language="zh-CN",
    )
    assert document.title == "闪光弹"
    assert document.source_uri == "upload://cs2-guide.md"
    assert "Micole" not in document.source_uri


def test_document_ids_are_content_stable() -> None:
    first = load_knowledge_bytes(b"same content", "a.txt")
    second = load_knowledge_bytes(b"same content", "a.txt")
    changed = load_knowledge_bytes(b"changed content", "a.txt")
    assert first.document_id == second.document_id
    assert first.document_id != changed.document_id


def test_docx_upload_preserves_headings_and_tables() -> None:
    from docx import Document

    buffer = BytesIO()
    source = Document()
    source.add_heading("CS2 经济管理", level=1)
    source.add_paragraph("购买结论必须有金钱和回合上下文。")
    table = source.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "证据"
    table.cell(0, 1).text = "结论"
    source.save(buffer)

    document = load_knowledge_bytes(buffer.getvalue(), "guide.docx")
    assert document.title == "CS2 经济管理"
    assert "# CS2 经济管理" in document.content
    assert "证据 | 结论" in document.content
    assert document.source_uri == "upload://guide.docx"


def test_markdown_chunking_preserves_heading_and_stable_ids() -> None:
    document = load_knowledge_bytes(
        ("# Positioning\n\n" + "Use cover before peeking. " * 80).encode(),
        "guide.md",
    )
    chunker = KnowledgeChunker(max_chars=260, overlap_chars=40)
    first = chunker.chunk(document)
    second = chunker.chunk(document)
    assert len(first) > 1
    assert [chunk.chunk_id for chunk in first] == [chunk.chunk_id for chunk in second]
    assert all(chunk.heading == "Positioning" for chunk in first)


def test_numbered_docx_sections_are_semantically_separated() -> None:
    document = KnowledgeDocument(
        document_id="manual",
        title="CS2 basic rules",
        source_uri="local://manual.docx",
        content=(
            "1. Victory Conditions\n\nCT wins when the planted bomb is defused.\n\n"
            "36.1 Economy Values\n\nversion_sensitive: true\nlast_verified: 2026-08-27\n"
            "Rifle kill reward is $300.\n\n"
            "76. Situation Decision Rule: 1v4 Post-Plant\n\n"
            "Consider saving an expensive weapon unless the clutch is realistic."
        ),
    )
    chunks = KnowledgeChunker(max_chars=400).chunk(document)
    assert [chunk.heading for chunk in chunks] == [
        "1. Victory Conditions",
        "36.1 Economy Values",
        "76. Situation Decision Rule: 1v4 Post-Plant",
    ]
    assert chunks[0].layer == KnowledgeLayer.GAME_RULES
    assert chunks[0].rule_strength == RuleStrength.HARD_RULE
    assert chunks[1].layer == KnowledgeLayer.DYNAMIC_GAME_DATA
    assert chunks[1].version_sensitive is True
    assert chunks[2].layer == KnowledgeLayer.SITUATION_DECISIONS
    assert chunks[2].rule_strength == RuleStrength.CONTEXTUAL_RECOMMENDATION


def test_situation_first_retrieval_reserves_a_decision_passage() -> None:
    stores = {layer: InMemoryKnowledgeStore(KeywordEmbedding()) for layer in KnowledgeLayer}
    store = LayeredKnowledgeStore(stores)
    store.upsert([
        KnowledgeChunk(
            chunk_id="weapon_awp",
            document_id="weapon",
            title="AWP",
            source_uri="builtin://weapon",
            content="AWP weapon economy",
            chunk_index=0,
            layer=KnowledgeLayer.TACTICAL_FUNDAMENTALS,
        ),
        KnowledgeChunk(
            chunk_id="save_1v4",
            document_id="situation",
            title="1v4 post-plant",
            source_uri="builtin://situation",
            content="1v4 post-plant AWP no kit save unless a clutch is realistic",
            chunk_index=0,
            layer=KnowledgeLayer.SITUATION_DECISIONS,
            rule_strength=RuleStrength.CONTEXTUAL_RECOMMENDATION,
        ),
    ])
    retriever = KnowledgeRetriever(store, min_score=-1.0)
    hits = retriever.retrieve_for_decision(
        "AWP post-plant decision",
        DecisionContext(
            side="ct",
            alive_teammates=1,
            alive_enemies=4,
            bomb_state="planted",
            weapon="awp",
            defuse_kit=False,
            time_remaining_sec=18,
        ),
        top_k=2,
    )
    assert hits[0].chunk.chunk_id == "save_1v4"
    assert {hit.chunk.chunk_id for hit in hits} == {"save_1v4", "weapon_awp"}


def test_dynamic_records_have_stable_ids_and_freshness_metadata() -> None:
    chunks = [
        chunk
        for document in dynamic_game_documents()
        for chunk in KnowledgeChunker(max_chars=500).chunk(document)
    ]
    assert chunks
    assert all(chunk.layer == KnowledgeLayer.DYNAMIC_GAME_DATA for chunk in chunks)
    assert all(chunk.version_sensitive for chunk in chunks)
    assert all(chunk.last_verified and chunk.source_urls for chunk in chunks)
    first_ids = [chunk.chunk_id for chunk in chunks]
    updated = dynamic_game_documents()[0].model_copy(update={
        "content": dynamic_game_documents()[0].content + "\n\nPatch note text changed.",
    })
    updated_ids = [
        chunk.chunk_id for chunk in KnowledgeChunker(max_chars=500).chunk(updated)
    ]
    assert updated_ids[0] == first_ids[0]


def test_in_memory_retrieval_returns_relevant_source() -> None:
    flash = load_knowledge_bytes(
        b"Flash exposure should be reviewed from cover.", "flash.txt",
    )
    economy = load_knowledge_bytes(
        b"Economy choices require verified money evidence.", "economy.txt",
    )
    store = InMemoryKnowledgeStore(KeywordEmbedding())
    result = index_documents([flash, economy], store)
    assert result.document_count == 2
    retriever = KnowledgeRetriever(store, min_score=0.01)
    hits = retriever.retrieve("flash cover", top_k=1)
    assert hits[0].chunk.source_uri == "upload://flash.txt"


def test_repo_policy_can_be_loaded_without_absolute_path_leak() -> None:
    with TemporaryDirectory() as tmp:
        path = Path(tmp) / "policy.md"
        path.write_text("# Policy\n\nEvidence first.", encoding="utf-8")
        document = load_knowledge_document(path)
    assert document.source_uri == "local://policy.md"
    assert tmp not in document.source_uri


def test_chroma_store_round_trip_when_optional_dependency_is_installed() -> None:
    import pytest

    pytest.importorskip("chromadb")
    # Chroma keeps its HNSW mmap open until process exit on Windows.
    with TemporaryDirectory(ignore_cleanup_errors=True) as tmp:
        store = ChromaKnowledgeStore(tmp, KeywordEmbedding())
        index_documents([
            load_knowledge_bytes(
                "flash cover 闪光 掩体".encode("utf-8"), "flash.txt",
            ),
        ], store)
        hits = store.query("flash cover", top_k=1)
    assert hits[0].chunk.source_uri == "upload://flash.txt"
    assert hits[0].score > 0.8
